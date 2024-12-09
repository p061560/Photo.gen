from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()

# 設定整份文件字型
style = document.styles['Normal']
font = style.font
font.name = '標楷體'
font.size = Pt(11)

# 第一行標題「現場勘察照片」
heading_para = document.add_paragraph()
heading_run = heading_para.add_run("現場勘察照片")
heading_run.bold = False
heading_run.font.name = '標楷體'
heading_run.font.size = Pt(20)
heading_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# 建立表格
table = document.add_table(rows=2, cols=8, style='Table Grid')
table.autofit = False

# 第一、第二列欄寬
first_and_second_row_widths = [Cm(1.17), Cm(1.45), Cm(0.75), Cm(1.57), Cm(1.88), Cm(2.04), Cm(1), Cm(6.6)]

# 設定第一、第二列的高度和欄寬
row1 = table.rows[0]
row2 = table.rows[1]
row1.height = Cm(0.71)
row2.height = Cm(0.71)
row1.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
row2.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

for i, width in enumerate(first_and_second_row_widths):
    table.columns[i].width = width

# 合併儲存格
table.cell(0, 0).merge(table.cell(1, 0))
table.cell(0, 6).merge(table.cell(1, 6))
table.cell(0, 7).merge(table.cell(1, 7))

# 設定儲存格內容
cell_texts = [
    ("案號", table.cell(0, 0)),
    ("年", table.cell(0, 1)),
    ("月", table.cell(0, 2)),
    ("警局", table.cell(0, 3)),
    ("分局", table.cell(0, 4)),
    ("案次編號", table.cell(0, 5)),
    ("案由", table.cell(0, 6)),
    ("", table.cell(0, 7))
]

for text, cell in cell_texts:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '標楷體'
    run.font.size = Pt(11)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 新增第三列
row3 = table.add_row()
row3.height = Cm(8.71)
row3.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
table.cell(2, 0).merge(table.cell(2, 7))

# 填入第三列的內容
merged_cell = table.cell(2, 0)
merged_cell.text = ""
p = merged_cell.paragraphs[0]
run = p.add_run("編號1")
run.font.bold = True
run.font.name = '標楷體'
run.font.size = Pt(11)

# 新增第四列
row4 = table.add_row()
row4.height = Cm(0.85)
row4.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

# 設定第四列的欄寬 [2.27, 3.72, 1.69, 4.8, 1.51, 2.47]
fourth_row_widths = [Cm(2.27), Cm(3.72), Cm(1.69), Cm(4.8), Cm(1.51), Cm(2.47)]

# 合併第六和第七欄
table.cell(3, 6).merge(table.cell(3, 7))

# 為第四列設定欄寬
for idx, width in enumerate(fourth_row_widths):
    if idx < 6:  # 因為只有6欄，防止索引超出
        cell = table.cell(3, idx)
        p = cell.paragraphs[0]

# 第6欄文字填入
merged_cell_last = table.cell(3, 6)
p_last = merged_cell_last.paragraphs[0]
run_last = p_last.add_run("第6欄合併")
run_last.font.name = '標楷體'
run_last.font.size = Pt(11)
p_last.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
merged_cell_last.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

document.save("output.docx")
